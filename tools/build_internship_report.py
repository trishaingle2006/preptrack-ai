from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "deliverables"
OUT_DIR.mkdir(exist_ok=True)
OUTPUT = OUT_DIR / "PrepTrack_AI_Internship_Report_Trisha_Ingle.docx"

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
PURPLE = "4F46E5"
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"
MUTED = "5F6B7A"
WHITE = "FFFFFF"
GREEN = "157347"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)


def set_run(run, size=11, color=NAVY, bold=False, italic=False, font="Calibri"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def configure_section(section, first=False):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = first


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])
    set_run(run, size=9, color=MUTED)


def add_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run("PREPTRACK AI  |  INTERNSHIP PROJECT REPORT"), size=8.5, color=MUTED, bold=True)
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    set_run(p.add_run("Trisha Ingle  |  "), size=9, color=MUTED)
    add_page_field(p)


def add_para(doc, text="", *, size=11, color=NAVY, bold=False, italic=False,
             align=None, before=0, after=6, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    p.paragraph_format.keep_with_next = keep
    if align is not None:
        p.alignment = align
    set_run(p.add_run(text), size=size, color=color, bold=bold, italic=italic)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    p.paragraph_format.keep_together = True
    p.add_run(text)
    return p


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(header), size=9.5, color=NAVY, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_run(p.add_run(str(value)), size=9.5, color=NAVY)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_callout(doc, label, value, fill=PALE_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    set_run(p.add_run(label.upper() + "\n"), size=9, color=BLUE, bold=True)
    set_run(p.add_run(value), size=11, color=NAVY, bold=True)
    set_table_geometry(table, [9360])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_page_break(doc):
    doc.add_page_break()


doc = Document()
configure_styles(doc)
section = doc.sections[0]
configure_section(section, first=True)
add_header_footer(section)

# Editorial cover
add_para(doc, "INTERNSHIP PROJECT REPORT", size=11, color=PURPLE, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, before=72, after=20)
add_para(doc, "PrepTrack AI", size=30, color=NAVY, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
add_para(doc, "AI-Powered Placement Preparation Platform", size=16, color=DARK_BLUE,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=24)
add_para(doc, "An integrated platform for adaptive interviews, recruiter simulation, placement readiness, peer challenges, practice, and secure role-based collaboration.",
         size=11.5, color=MUTED, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=52)
add_callout(doc, "Prepared by", "Trisha Ingle\nFrontend Developer and Computer Science Student")
add_para(doc, "Internship platform: ElevanceSkills\nReport year: 2026", size=10.5, color=MUTED,
         align=WD_ALIGN_PARAGRAPH.CENTER, before=28, after=12)
add_para(doc, "Live application: https://preptrack-ai-sepia.vercel.app\nSource repository: https://github.com/trishaingle2006/preptrack-ai",
         size=9.5, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, before=26, after=0)

add_page_break(doc)
add_heading(doc, "Project at a Glance", 1)
add_para(doc, "PrepTrack AI is a production-deployed web application that consolidates six internship assignments into one cohesive placement-preparation workspace. The solution began as a single-file college project and was redesigned into a modular React architecture with Firebase authentication, Cloud Firestore, Vercel serverless functions, and Gemini-assisted evaluation.")
add_table(doc, ["Area", "Implementation"], [
    ("Frontend", "React, Vite, responsive CSS, route-based workspaces"),
    ("Backend", "Vercel serverless API with server-authoritative operations"),
    ("Data and identity", "Firebase Authentication and Cloud Firestore"),
    ("AI capability", "Gemini-assisted interview, recruiter, readiness, and challenge evaluation"),
    ("Security", "Email verification, session tracking, alerts, RBAC, and secure admin/mentor APIs"),
    ("Deployment", "Vercel production hosting; Firebase Spark plan for authentication and Firestore"),
], [2700, 6660])
add_heading(doc, "Executive Summary", 1)
add_para(doc, "The project addresses a common placement-preparation problem: students often use disconnected resources for aptitude practice, interview preparation, resume review, and peer competition. PrepTrack AI brings these activities into a single evidence-driven platform. Performance data flows across modules so that practice completion, interview history, readiness assessments, and challenge results can contribute to a clearer picture of candidate progress.")
add_para(doc, "The final system is publicly accessible, backed by a documented GitHub repository, and tested with real user flows. Verified outcomes include an adaptive interview that progressed from easy to hard, an Amazon recruiter simulation score of 9.4/10, readiness improvement from 69 to 71, a peer-challenge score of 95 with rank #1, and working student, mentor, and administrator permissions.")
add_callout(doc, "Submission assets", "Hosted URL, GitHub repository, project documentation, daily work log, evidence records, and this internship report are prepared for final review.", fill="F4F6F9")

add_page_break(doc)
add_heading(doc, "1. Introduction", 1)
add_para(doc, "Placement preparation requires more than memorizing answers. Candidates must build foundations, practise communication, respond to changing question difficulty, understand company expectations, evaluate their overall readiness, and maintain consistent effort. PrepTrack AI was developed to support this complete preparation cycle through one secure, responsive web application.")
add_para(doc, "The project combines structured practice with AI-assisted evaluation. It provides immediate feedback while retaining objective evidence such as completed questions, interview scores, difficulty progression, readiness reports, leaderboard points, and mentor comments. The product is educational: company simulation profiles represent broad interview patterns and are not presented as official hiring processes.")
add_heading(doc, "2. Background", 1)
add_para(doc, "The original PrepTrack prototype was a large single HTML file containing styling, questions, authentication calls, progress tracking, and notes. Although functional as a college project, this structure was difficult to maintain and extend. User data depended heavily on browser storage, several actions were directly attached to HTML elements, and the application did not have a secure backend for AI or role-management operations.")
add_para(doc, "The internship scope required multiple advanced assignments. Instead of treating them as unrelated demonstrations, the project was re-architected as a unified product. Shared authentication, user profiles, Firestore records, reusable layouts, secure API calls, and consistent design patterns now connect every module.")
add_heading(doc, "Problem Statement", 2)
add_para(doc, "Build an accessible placement-preparation platform that can adapt interview difficulty, simulate recruiter expectations, synthesize readiness evidence, support competitive peer practice, and protect account and role data—while remaining deployable on free service tiers.")

add_page_break(doc)
add_heading(doc, "3. Learning Objectives", 1)
add_para(doc, "The internship project was designed around the following learning objectives:")
for item in [
    "Transform a monolithic prototype into a modular, maintainable React application.",
    "Design adaptive workflows that use answer quality to control interview difficulty and follow-up questions.",
    "Integrate generative AI safely through a server-side API with validation, rate limiting, and local fallback behavior.",
    "Model user progress, interviews, challenges, readiness history, sessions, and security events in Firestore.",
    "Implement authentication and authorization that distinguish student, mentor, and administrator responsibilities.",
    "Deploy, diagnose, and verify a production application using Firebase, Vercel, browser testing, logs, builds, and linting.",
    "Document the project clearly through a README, requirements checklist, project plan, daily log, source repository, and final report.",
]:
    add_bullet(doc, item)
add_heading(doc, "Engineering Principles Applied", 2)
add_table(doc, ["Principle", "Application in PrepTrack AI"], [
    ("Modularity", "Pages, data sources, security permissions, and service calls are separated by responsibility."),
    ("Progressive enhancement", "Local evaluation preserves usability when the external AI service is unavailable."),
    ("Least privilege", "Routes and server operations check role permissions before returning protected data."),
    ("Evidence over assumption", "Scores, history, and activity records support reports and mentor review."),
    ("Responsive design", "Desktop sidebar and mobile navigation provide access across screen sizes."),
], [2300, 7060])

add_page_break(doc)
add_heading(doc, "4. Activities and Tasks", 1)
add_heading(doc, "4.1 Foundation and Migration", 2)
add_para(doc, "The first phase migrated the existing PrepTrack prototype into a Vite and React application. The work separated routing, pages, data, services, authentication context, permission definitions, and styling. Practice questions and notes were preserved while dashboards were converted into dynamic views that read current Firestore-backed evidence.")
add_heading(doc, "4.2 Core Feature Development", 2)
add_table(doc, ["Module", "Key activities completed"], [
    ("Adaptive Interview Engine", "Created question banks, evaluation flow, answer repetition detection, skip handling, dynamic difficulty, contextual follow-ups, duplicate prevention, and final progression reports."),
    ("Recruiter Simulator", "Built educational company profiles, target-role sessions, configurable standards, adaptive questioning, weighted evaluation emphasis, and company-specific reports."),
    ("Placement Readiness Engine", "Combined resume content, practice completion, and interview evidence into scores, classifications, history, improvement areas, and immediate next steps."),
    ("Peer Challenge Arena", "Implemented daily and weekly challenges, final submissions, AI scoring, points, leaderboards, ranks, streaks, and badges."),
    ("Enterprise Authentication", "Added registration, email verification, password reset and strength rules, session replacement, login history, alerts, account status, and role checks."),
    ("Mentor and Admin Workspaces", "Implemented secure user listing, role/status updates, platform settings, student evidence review, and structured mentor feedback."),
], [2400, 6960])
add_heading(doc, "4.3 Deployment and Documentation", 2)
add_para(doc, "The application was built and deployed to Vercel with environment variables stored as sensitive production values. Firebase Authentication and Firestore remained on the Spark no-cost plan. Firestore rules and indexes were deployed separately. The source was committed to Git, pushed to GitHub, and accompanied by setup instructions and project documentation.")

add_page_break(doc)
add_heading(doc, "5. Skills and Competencies Developed", 1)
add_heading(doc, "Frontend Engineering", 2)
for item in [
    "React component architecture, hooks, context, routing, forms, and stateful multi-step workflows.",
    "Responsive interface design using CSS Grid, Flexbox, desktop navigation, and mobile navigation.",
    "Reusable data-driven cards, reports, tables, status states, and permission-aware navigation.",
    "Production build diagnosis, lint remediation, bundle awareness, and browser verification.",
]: add_bullet(doc, item)
add_heading(doc, "Backend, Data, and AI", 2)
for item in [
    "Firebase Authentication lifecycle and Cloud Firestore document/subcollection modeling.",
    "Serverless request handling, Firebase Admin SDK usage, input validation, structured errors, and rate limiting.",
    "Prompt-based evaluation with Gemini, JSON response handling, fallbacks, and service failure diagnosis.",
    "Server-authoritative admin and mentor operations that avoid trusting client-side roles.",
]: add_bullet(doc, item)
add_heading(doc, "Professional Competencies", 2)
for item in [
    "Breaking a large objective into testable implementation phases.",
    "Debugging through error messages, logs, build output, and direct production testing.",
    "Making cost-aware architecture decisions while preserving project requirements.",
    "Maintaining project evidence through Git, GitHub, reports, plans, checklists, and daily updates.",
]: add_bullet(doc, item)

add_page_break(doc)
add_heading(doc, "6. Feedback and Evidence", 1)
add_para(doc, "The project was verified through production user flows rather than relying only on code completion. The following evidence was recorded during testing:")
add_table(doc, ["Test area", "Observed evidence", "Result"], [
    ("Adaptive Interview", "Five answered questions; difficulty progressed easy → medium → hard", "10/10 average"),
    ("Recruiter Simulator", "Amazon Frontend Developer simulation; five questions", "9.4/10 vs 7.5 standard"),
    ("Placement Readiness", "Report regenerated after completing three practice questions", "69 → 71"),
    ("Challenge Arena", "Daily HR challenge evaluated and persisted", "95/100; rank #1; streak 1"),
    ("Account Security", "Verified email, replaced prior session, recorded login activity and alert", "Verified"),
    ("Role permissions", "Student blocked; mentor review and admin management tested", "Verified"),
    ("Production delivery", "Vite build, Vercel deploy, Firestore rules deploy, GitHub push", "Successful"),
], [2100, 5060, 2200])
add_heading(doc, "Mentor Feedback Evidence", 2)
add_callout(doc, "Recorded feedback", "Excellent interview performance and strong improvement in placement readiness. Continue practising timed coding assessments and strengthen TypeScript and automated testing skills.", fill="F4F6F9")
add_heading(doc, "Interpretation", 2)
add_para(doc, "The evidence shows that the platform is not a static interface. Interview difficulty changes in response to performance, readiness reflects new activity, challenge completion updates the leaderboard, and protected workspaces expose different capabilities by role. The feedback also identifies meaningful next-learning priorities rather than claiming that the project eliminates the need for continued development.")

add_page_break(doc)
add_heading(doc, "7. Challenges and Solutions", 1)
add_table(doc, ["Challenge", "Resolution"], [
    ("Monolithic legacy structure", "Reorganized the single-file project into React pages, context, data modules, services, permissions, and shared styling."),
    ("Firebase Cloud Functions required a Blaze plan", "Kept Firebase on Spark and deployed the secure API as Vercel serverless functions, preserving a free architecture."),
    ("Gemini requests returned 404/model errors", "Diagnosed production logs, corrected the model/API integration, redeployed, and retained local evaluation as a resilience path."),
    ("SPA routes returned 404 after browser refresh", "Added Vercel routing configuration so application routes rewrite to the React entry point while API traffic remains separate."),
    ("Client Firestore access caused permission errors", "Moved privileged admin and mentor operations to server-authoritative APIs with role revalidation."),
    ("Session and verification state initially appeared empty", "Completed email verification, refreshed authentication state, and tested sequential sign-ins to validate session replacement and alerts."),
    ("Build and lint errors interrupted deployment", "Fixed invalid asynchronous code, JSX syntax, unused variables, and hook/export structure before production releases."),
    ("Git was unavailable and repository ownership differed", "Installed Git, configured the workspace as a safe directory, staged only intended files, committed, authenticated, and pushed main."),
], [3200, 6160])
add_heading(doc, "Key Lesson", 2)
add_para(doc, "The most important lesson was that production engineering includes far more than feature coding. Authentication state, service plans, environment variables, browser routing, authorization boundaries, logging, data rules, version control, and documentation all determine whether an application is genuinely usable and defensible.")

add_page_break(doc)
add_heading(doc, "8. Outcomes and Impact", 1)
add_para(doc, "The internship produced a complete, publicly accessible placement-preparation platform rather than six disconnected prototypes. The system gives students a continuous workflow: practise fundamentals, attempt adaptive interviews, experience recruiter-style simulations, generate a readiness roadmap, join peer challenges, and receive mentor feedback within a protected account.")
add_heading(doc, "Measured Project Outcomes", 2)
for item in [
    "One integrated application containing six core internship modules plus practice, notes, mentor, admin, and security workspaces.",
    "Production deployment at a stable Vercel URL with Firebase Authentication and Firestore on the no-cost Spark plan.",
    "AI-assisted evaluation working in production, with local fallback behavior for service interruptions.",
    "Server-enforced role management for student, mentor, and administrator workflows.",
    "Persistent evidence across interviews, readiness reports, practice completion, challenge ranking, sessions, alerts, and mentor feedback.",
    "Published GitHub repository with setup documentation, requirements tracking, project planning, and work logs.",
]: add_bullet(doc, item)
add_heading(doc, "Value to Learners", 2)
add_para(doc, "PrepTrack AI makes progress visible. Instead of receiving only a final score, candidates can see difficulty progression, strong and weak concepts, readiness components, improvement roadmaps, rankings, streaks, and mentor feedback. This encourages repeated, focused practice and provides evidence that can be discussed during placement preparation.")
add_heading(doc, "Future Improvement Opportunities", 2)
add_para(doc, "Recommended future work includes TypeScript migration, automated unit and end-to-end tests, code splitting to reduce the main JavaScript bundle, richer accessibility audits, expanded question banks, improved analytics, and additional safeguards for production-scale AI usage.")

add_page_break(doc)
add_heading(doc, "9. Conclusion", 1)
add_para(doc, "PrepTrack AI successfully evolved from a basic single-file practice portal into a modular, secure, AI-assisted placement-preparation platform. The project demonstrates frontend development, cloud data integration, serverless APIs, authentication, role-based authorization, adaptive evaluation, production deployment, debugging, and technical documentation in one coherent application.")
add_para(doc, "The work also strengthened independent problem solving and ownership. Several obstacles—including service-plan restrictions, AI model errors, route refresh failures, Firestore permissions, session behavior, build failures, and Git setup—were resolved through evidence-based diagnosis and iterative testing. The final result is a functioning product with traceable outcomes, not only a design concept.")
add_para(doc, "The project meets its central objective: providing students with an integrated environment to build confidence, prove readiness, and prepare more intentionally for placements. It also establishes a strong foundation for continued learning in TypeScript, testing, scalable state management, performance optimization, and CI/CD.")
add_callout(doc, "Final status", "Core development and production validation are complete. The remaining action is the final one-time internship portal submission after reviewing the hosted URL, GitHub repository, and report file.", fill="EAF6EF")

add_heading(doc, "Appendix A — Submission Links", 1)
add_table(doc, ["Deliverable", "Location"], [
    ("Hosted application", "https://preptrack-ai-sepia.vercel.app"),
    ("GitHub repository", "https://github.com/trishaingle2006/preptrack-ai"),
    ("Project README", "Available in the GitHub repository"),
    ("Daily work log", "DAILY_LOG.md in the GitHub repository"),
    ("Requirements checklist", "REQUIREMENTS_CHECKLIST.md in the GitHub repository"),
], [2600, 6760])

for sec in doc.sections:
    configure_section(sec, first=(sec is doc.sections[0]))

doc.core_properties.title = "PrepTrack AI Internship Project Report"
doc.core_properties.subject = "AI-Powered Placement Preparation Platform"
doc.core_properties.author = "Trisha Ingle"
doc.core_properties.keywords = "PrepTrack AI, internship, React, Firebase, Gemini AI, placement preparation"
doc.save(OUTPUT)
print(OUTPUT)
